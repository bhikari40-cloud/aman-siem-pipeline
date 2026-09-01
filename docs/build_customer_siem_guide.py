"""Generates the Aman SIEM Pipeline customer-facing SIEM configuration guide.

Scope, deliberately narrow (Brian's correction, 2026-09-01): this describes
ONE thing -- the one pipeline that is actually running (ecs_syslog_webhook.py)
and how to configure a SIEM to receive exactly what it sends. No mention of
the second, not-currently-running engine, no per-SIEM "status" tiers -- that
distinction is real but is not something a customer configuring their SIEM
needs to know.

Matches the visual style already established in
docs/Aman-SIEM-Pipeline-Architecture-and-Backend-Setup-Guide.docx (Calibri,
navy/blue/slate palette, SF Mono for code, Table Grid tables).
"""
from __future__ import annotations

import pathlib

import docx
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1A, 0x27, 0x33)
BLUE = RGBColor(0x2A, 0x52, 0xB8)
SLATE = RGBColor(0x47, 0x55, 0x69)
CODE_BLACK = RGBColor(0x0A, 0x0A, 0x0A)
RED_TEXT = RGBColor(0xB4, 0x23, 0x18)

doc = docx.Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def title_block() -> None:
    p = doc.add_paragraph()
    r = p.add_run("IntelliBroń Aman")
    r.font.size = Pt(16)
    r.font.color.rgb = SLATE

    p = doc.add_paragraph()
    r = p.add_run("SIEM Configuration Guide")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    r = p.add_run("What Aman sends, and how to set your SIEM up to receive it")
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    doc.add_paragraph()

    def meta(label: str, value: str, value_color=SLATE) -> None:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r2 = p.add_run(value)
        r2.font.color.rgb = value_color

    meta("Status", "DRAFT — for internal review before this goes to any customer", RED_TEXT)
    meta("Prepared by", "Brian Hikari Janna — Cyber Security Product Researcher, ITSEC Asia")
    meta("Date", "2026-09-01")
    meta(
        "Source repo",
        "github.com/bhikari40-cloud/aman-siem-pipeline, branch customer-onboarding-gate",
    )
    doc.add_paragraph()


def h1(text: str) -> None:
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(21)
    r.font.color.rgb = NAVY
    r.font.bold = True


def h2(text: str) -> None:
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = BLUE
    r.font.bold = True


def body(text: str) -> None:
    doc.add_paragraph(text)


def numbered(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "SF Mono"
    r.font.size = Pt(9.5)
    r.font.color.rgb = CODE_BLACK


def fields_table(rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        c0 = row.cells[0]
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.color.rgb = NAVY
        r0.font.size = Pt(10.5)
        c1 = row.cells[1]
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.name = "SF Mono" if any(ch in value for ch in ("/", "://", "<", "{")) else "Calibri"
        r1.font.size = Pt(10.5)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
title_block()

h1("1. What Aman sends")
body(
    "Aman delivers each blocked DNS event as a single HTTPS POST request to a URL you provide. "
    "There is no batching and no separate “native alert” — each blocked query is sent to you "
    "as soon as it happens."
)
body(
    "The body of each request is one line of plain text: a standard syslog (RFC 5424) header, "
    "followed by a JSON object holding the event's fields — domain, source IP, severity, and "
    "so on — using Elastic Common Schema (ECS) field names. This is a real, working example:"
)
code(
    '<34>1 2026-09-01T02:14:07.000Z aman-pipeline aman-dns - - - '
    '{"@timestamp":"2026-09-01T02:14:07.000Z","event":{"type":["denied"],"category":["network"],'
    '"action":"blocked"},"source":{"ip":"203.153.118.242"},"dns":{"question":{"name":'
    '"phishing-kit.example","type":"ANY"}},"severity":"Critical"}'
)
body(
    "The part before the JSON is the syslog header; everything from the first { onward is the "
    "JSON object."
)
body(
    "Every request also carries an Authorization header, using a token you provide: "
    "Authorization: Bearer <your token>. Splunk is the one exception — see the Splunk section "
    "below for the login format it requires instead."
)

h1("2. What to give Aman")
fields_table([
    ("Webhook URL", "the HTTPS address Aman should send each alert to"),
    ("Auth token", "a token you generate — Aman sends it back to you in every request"),
])

h1("3. Setting up Splunk")
body(
    "Splunk has a built-in way to receive exactly this kind of data: its HTTP Event Collector "
    "(HEC)."
)
numbered([
    "In Splunk Web: Settings → Data Inputs → HTTP Event Collector → New Token.",
    "Name it (for example, “IntelliBroń Aman”), choose or create a source type, and continue.",
    "Choose which index should receive the data, then Review and Submit.",
    "Under Settings → Data Inputs → HTTP Event Collector → Global Settings, confirm "
    "“All Tokens” is enabled, and enable SSL if it is not already on.",
    "Copy the token value Splunk generates — you will not be able to view it again later.",
])
fields_table([
    ("Webhook URL", "https://<your-splunk-host>:8088/services/collector/raw"),
    ("Auth token", "the HEC token from step 5 above"),
])
body(
    "Use the /services/collector/raw path exactly as shown, not the bare /services/collector "
    "address — that address only accepts JSON, and will reject this format."
)
body(
    "Splunk indexes each line as raw text by default — it does not automatically split the JSON "
    "portion into separate searchable fields. To search on individual fields (domain, severity, "
    "and so on), use Splunk's spath command at search time, or add field extraction "
    "(KV_MODE=json) to the sourcetype you chose in step 2."
)

h1("4. Setting up other SIEMs")
body(
    "The platforms below can each receive this exact webhook directly. Two platforms — Wazuh "
    "and Microsoft Sentinel — do not currently have a way to accept this format directly; "
    "those sections explain the practical alternative."
)

h2("Elastic (via Logstash)")
body(
    "Elasticsearch's own Bulk API requires a specific JSON format and cannot accept this. "
    "Logstash's http input can, using its default plain-text mode."
)
numbered([
    "If you don't already run Logstash, install it alongside your Elastic Stack.",
    "Configure a pipeline with an http input (default codec is already plain text — no change "
    "needed) listening on a port of your choice, with TLS enabled, and an elasticsearch output "
    "pointing at your cluster.",
])
fields_table([("Webhook URL", "https://<your-logstash-host>:<port>/")])
body(
    "Logstash's http input does not check any credential by default. If you want to verify the "
    "token Aman sends, add a filter that checks the Authorization header value — or rely on "
    "network-level access control (firewall or VPN) instead. Note: Elastic's separate "
    "“Custom HTTP Endpoint” integration will not work for this — it requires a JSON body and "
    "rejects plain text."
)

h2("Graylog")
body("Graylog's GELF HTTP input requires a specific JSON structure and cannot accept this. Its separate Raw HTTP input can.")
numbered([
    "System → Inputs → select “Raw HTTP Input” → Launch new input.",
    "Set Bind Address and Port (enable TLS if you want HTTPS).",
    "Set “Authorization Header Value” to exactly Bearer, followed by a space, then the token "
    "you give Aman — for example: Bearer a1b2c3d4. This is what makes Graylog check the "
    "token instead of accepting requests from anyone who finds the URL.",
    "Save.",
])
fields_table([("Webhook URL", "https://<your-graylog-host>:<port>/raw")])

h2("Sumo Logic")
numbered([
    "Create a Hosted Collector with an HTTP Logs and Metrics Source.",
    "Copy the unique URL Sumo Logic generates for that source.",
])
fields_table([("Webhook URL", "the URL from step 2, shaped like https://<endpoint>/receiver/v1/http/<code>")])
body(
    "That URL is itself the credential — treat it like a password. Sumo Logic does not check "
    "the separate auth token Aman asks for, so any placeholder value there is fine."
)

h2("CrowdStrike Falcon LogScale")
numbered(["Create an ingest token in LogScale, scoped to the repository this data should land in."])
fields_table([
    ("Webhook URL", "https://<your-logscale-host>/api/v1/ingest/hec/raw"),
    ("Auth token", "the ingest token from step 1"),
])
body("This uses the same login format Aman sends by default, so no extra configuration is needed on Aman's side.")

h2("Datadog")
numbered(["In Datadog: Organization Settings → API Keys → New Key. Name it and copy the value."])
fields_table([
    ("Webhook URL", "https://http-intake.logs.datadoghq.com/api/v2/logs"),
    ("Auth token", "the API key from step 1"),
])
body(
    "Datadog reads the credential from a header named DD-API-KEY, not the more common "
    "Authorization header — Aman sends it correctly either way, this is just how Datadog itself "
    "expects it, in case you're checking request logs on your side."
)

h2("Wazuh — needs a relay you build")
body(
    "Wazuh's real alerting path only accepts syslog over TCP or UDP, not HTTP, and its "
    "indexer's API requires a different, specific JSON format. Neither can accept this webhook "
    "directly. If you need this integration, you would stand up a small receiver of your own "
    "that accepts Aman's webhook and forwards the data into your Wazuh manager over syslog. "
    "Contact your Aman representative to discuss this."
)

h2("Microsoft Sentinel — needs a relay you build")
body(
    "Sentinel's ingestion API requires a specific JSON structure matching a table schema you "
    "define, and cannot accept this format directly. The standard way around this is a small "
    "Azure Function or Logic App with an HTTP trigger that receives Aman's webhook, wraps it as "
    "JSON, and forwards it into Sentinel via that ingestion API. Contact your Aman "
    "representative if you'd like help scoping this."
)

h2("Anything else")
body(
    "Any endpoint that can accept an HTTPS POST request and read a plain-text body can receive "
    "this webhook, even if it isn't named above — check your platform's documentation for terms "
    "like “HTTP input,” “webhook,” or “custom log source.”"
)

OUTPUT_PATH = pathlib.Path(__file__).parent / "Aman-SIEM-Pipeline-Customer-SIEM-Configuration-Guide.docx"
doc.save(str(OUTPUT_PATH))
print(f"saved to {OUTPUT_PATH}")
